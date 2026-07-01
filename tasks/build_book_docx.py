"""Build the final formatted Word document for the SemanticGuard AI thesis."""

from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "SemanticGuard-AI-Book.md"
OUT = ROOT / "docs" / "SemanticGuard-AI-Book-Final-Formatted.docx"
IMAGES = ROOT / "docs" / "images"

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"
MAX_IMAGE_WIDTH = Inches(6.0)
MAX_IMAGE_HEIGHT = Inches(8.4)

HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)$")
CAPTION_RE = re.compile(r"^\*\*(Figure|Table)\s+(\d+):\s*(.*?)\*\*$")
IMAGE_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)$")
BULLET_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
NUMBER_RE = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
COMMENT_START_RE = re.compile(r"^\s*<!--")
COMMENT_END_RE = re.compile(r"-->\s*$")


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    root = run._r

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    root.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    root.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    root.append(separate)

    if placeholder:
        text = OxmlElement("w:t")
        text.text = placeholder
        root.append(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    root.append(end)
    return run


def set_run_font(run, *, size: int = 12, bold: bool = False, italic: bool = False, font: str = BODY_FONT) -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    properties = run._element.get_or_add_rPr()
    fonts = properties.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.append(fonts)
    for attribute in ("w:ascii", "w:hAnsi", "w:cs"):
        fonts.set(qn(attribute), font)


def clean_text(text: str) -> str:
    return html.unescape(text).replace("\xa0", " ").strip()


def add_inline(paragraph, text: str, *, size: int = 12, bold: bool = False, centered: bool = False) -> None:
    text = clean_text(text)
    if centered:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    token = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    for part in token.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size - 1, font=MONO_FONT)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, bold=bold)


def configure_section(section, *, page_format: str, start: int | None = None, first_page_blank: bool = False) -> None:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.different_first_page_header_footer = first_page_blank

    sect_pr = section._sectPr
    page_numbers = sect_pr.find(qn("w:pgNumType"))
    if page_numbers is None:
        page_numbers = OxmlElement("w:pgNumType")
        sect_pr.append(page_numbers)
    page_numbers.set(qn("w:fmt"), page_format)
    if start is not None:
        page_numbers.set(qn("w:start"), str(start))

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.text = ""
    add_field(paragraph, "PAGE")
    for run in paragraph.runs:
        set_run_font(run, size=11)

    if first_page_blank:
        first_footer = section.first_page_footer
        first_footer.is_linked_to_previous = False
        first_footer.paragraphs[0].text = ""


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal.element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.element.rPr.rFonts.set(qn("w:cs"), BODY_FONT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_specs = {
        "Heading 1": (16, WD_ALIGN_PARAGRAPH.CENTER),
        "Heading 2": (14, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 3": (12, WD_ALIGN_PARAGRAPH.LEFT),
    }
    for name, (size, alignment) in heading_specs.items():
        style = styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style.element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.paragraph_format.alignment = alignment
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    try:
        caption = styles["Caption"]
    except KeyError:
        caption = styles.add_style("Caption", 1)
    caption.font.name = BODY_FONT
    caption.font.size = Pt(11)
    caption.font.bold = True
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_image(doc: Document, image_path: Path, *, centered: bool = True) -> None:
    if not image_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(paragraph, f"[Missing image: {image_path.name}]", size=11)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run()
    picture = run.add_picture(str(image_path), width=MAX_IMAGE_WIDTH)
    if picture.height > MAX_IMAGE_HEIGHT:
        ratio = MAX_IMAGE_HEIGHT / picture.height
        picture.height = int(picture.height * ratio)
        picture.width = int(picture.width * ratio)


def add_caption(doc: Document, label: str, number: str, title: str) -> None:
    """Write a caption that uses a Word SEQ field so the List of Figures / List
    of Tables fields can collect it automatically."""
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = paragraph.add_run(f"{label} ")
    set_run_font(label_run, size=11, bold=True)
    seq_run = add_field(paragraph, f"SEQ {label} \\* ARABIC", number)
    set_run_font(seq_run, size=11, bold=True)
    title_run = paragraph.add_run(f": {clean_text(title)}")
    set_run_font(title_run, size=11, bold=True)


def add_toc_field(doc: Document, instruction: str, placeholder: str) -> None:
    """Insert an auto-updating Word table-of-contents style field."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    field_run = add_field(paragraph, instruction, placeholder)
    set_run_font(field_run, size=12)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip().strip("|")
        cells = [clean_text(cell) for cell in raw.split("|")]
        if not all(set(cell) <= set("-: ") for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for col_index in range(columns):
        text = rows[0][col_index] if col_index < len(rows[0]) else ""
        cell = table.rows[0].cells[col_index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(paragraph, text, bold=True)

    for row in rows[1:]:
        cells = table.add_row().cells
        for col_index in range(columns):
            text = row[col_index] if col_index < len(row) else ""
            cell = cells[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            add_inline(paragraph, text, size=10)

    doc.add_paragraph()


def add_list_items(doc: Document, items: list[tuple[str, str, str]]) -> None:
    for kind, number, text in items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.45)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        prefix = f"{number}. " if kind == "number" else "- "
        run = paragraph.add_run(prefix)
        set_run_font(run)
        add_inline(paragraph, text)


def build_document() -> None:
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0], page_format="lowerRoman", first_page_blank=True)

    in_comment = False
    in_center_block = False
    body_section_started = False
    pending_page_break = False
    list_items: list[tuple[str, str, str]] = []

    def flush_list() -> None:
        nonlocal list_items
        if list_items:
            add_list_items(doc, list_items)
            list_items = []

    index = 0
    skip_static_list = False
    auto_list_fields = {
        "TABLE OF CONTENTS": ('TOC \\o "1-3" \\h \\z \\u', "Right-click and choose 'Update Field' to build the table of contents."),
        "LIST OF TABLES": ('TOC \\h \\z \\c "Table"', "Right-click and choose 'Update Field' to build the list of tables."),
        "LIST OF FIGURES": ('TOC \\h \\z \\c "Figure"', "Right-click and choose 'Update Field' to build the list of figures."),
    }
    while index < len(lines):
        raw_line = lines[index]
        stripped = raw_line.strip()

        if in_comment:
            if COMMENT_END_RE.search(stripped):
                in_comment = False
            index += 1
            continue
        if COMMENT_START_RE.match(stripped):
            in_comment = not COMMENT_END_RE.search(stripped)
            index += 1
            continue

        if stripped.startswith("<div align=\"center\""):
            in_center_block = True
            index += 1
            continue
        if stripped == "</div>":
            in_center_block = False
            index += 1
            continue
        if stripped.startswith("<div") and "page-break-after" in stripped:
            flush_list()
            pending_page_break = True
            index += 1
            continue
        if stripped == "---":
            index += 1
            continue

        if not stripped:
            flush_list()
            index += 1
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            flush_list()
            skip_static_list = False
            level = len(heading_match.group(1))
            title = clean_text(heading_match.group(2))
            if title == "CHAPTER ONE" and not body_section_started:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, page_format="decimal", start=1)
                body_section_started = True
                pending_page_break = False
            elif pending_page_break:
                add_page_break(doc)
                pending_page_break = False
            paragraph = doc.add_heading(title.upper() if level == 1 else title, level=min(level, 3))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 or in_center_block else WD_ALIGN_PARAGRAPH.LEFT
            if title in auto_list_fields:
                instruction, placeholder = auto_list_fields[title]
                add_toc_field(doc, instruction, placeholder)
                skip_static_list = True
            index += 1
            continue

        if skip_static_list:
            index += 1
            continue

        if pending_page_break:
            flush_list()
            add_page_break(doc)
            pending_page_break = False

        caption_match = CAPTION_RE.match(stripped)
        if caption_match:
            flush_list()
            label, number, title = caption_match.groups()
            next_line = lines[index + 1].strip() if index + 1 < len(lines) else ""
            image_match = IMAGE_RE.match(next_line)
            if label == "Figure" and image_match:
                image_path = (SRC.parent / image_match.group(2)).resolve()
                add_image(doc, image_path)
                add_caption(doc, label, number, title)
                index += 2
                continue
            add_caption(doc, label, number, title)
            index += 1
            continue

        image_match = IMAGE_RE.match(stripped)
        if image_match:
            flush_list()
            image_path = (SRC.parent / image_match.group(2)).resolve()
            add_image(doc, image_path)
            index += 1
            continue

        if stripped.startswith("|"):
            flush_list()
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue

        number_match = NUMBER_RE.match(raw_line)
        if number_match:
            list_items.append(("number", number_match.group(2), number_match.group(3).strip()))
            index += 1
            continue

        bullet_match = BULLET_RE.match(raw_line)
        if bullet_match:
            list_items.append(("bullet", "", bullet_match.group(2).strip()))
            index += 1
            continue

        flush_list()
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if in_center_block else WD_ALIGN_PARAGRAPH.JUSTIFY
        if body_section_started and doc.paragraphs and doc.paragraphs[-1].text.startswith("http"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline(paragraph, stripped, centered=in_center_block)
        index += 1

    flush_list()
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build_document()