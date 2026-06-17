"""
Build a professionally formatted Word (.docx) version of the SemanticGuard AI
Product & Technical Proposal from the Markdown source.

Features:
  * Times New Roman 12, 1.5 line spacing, justified body text.
  * Title page, then auto-updating Table of Contents, List of Figures and
    List of Tables (Word fields - select all + F9 in Word to populate).
  * Every chapter (H1) starts on a new page.
  * Each major diagram is placed on its own page.
  * Figure and table captions are centred and auto-numbered via SEQ fields.

Usage:  py tasks/build_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "SemanticGuard-AI-Proposal.md"
OUT = ROOT / "docs" / "SemanticGuard-AI-Proposal-Formatted.docx"
IMAGES = ROOT / "docs" / "images"

# Map a figure's caption title to its image file in docs/images.
FIGURE_IMAGES = {
    "System Architecture": "System Architecture.png",
    "Use-Case Diagram": "Use Case Diagram.png",
    "Entity–Relationship Diagram": "ERD.png",
    "Assessment-Monitoring Sequence Diagram": "Sequence Diagram.png",
    "Work Breakdown Structure": "Work Breakdown Structure.png",
    "Authentication and Security Flow": "SYSTEM-LEVEL.png",
    "Gantt Chart": "GANTT CHART.png",
}

# Usable area on A4 with 1-inch margins.
MAX_IMG_W = Inches(6.2)
MAX_IMG_H = Inches(8.6)

BODY = "Times New Roman"
MONO = "Consolas"

CAPTION_RE = re.compile(r"^\*\*(Figure|Table)\s+\d+:\s*(.*?)\*\*$")
HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)$")
BULLET_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
NUMBER_RE = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")


# --------------------------------------------------------------------------- #
#  Low-level helpers
# --------------------------------------------------------------------------- #
def add_field(paragraph, instruction: str, placeholder: str = ""):
    """Insert a Word field (TOC, SEQ, ...) into a paragraph."""
    run = paragraph.add_run()
    r = run._r

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    r.append(instr)

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r.append(sep)

    if placeholder:
        t = OxmlElement("w:t")
        t.text = placeholder
        r.append(t)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r.append(end)
    return run


def set_run(run, *, font=BODY, size=12, bold=False, italic=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # ensure east-asian / hint also uses the same font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), font)


def add_inline(paragraph, text: str, *, base_size=12, base_bold=False):
    """Render **bold**, *italic* and `code` inline spans into runs."""
    token = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    for part in token.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            set_run(r, size=base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            set_run(r, font=MONO, size=base_size - 1)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = paragraph.add_run(part[1:-1])
            set_run(r, size=base_size, italic=True)
        else:
            r = paragraph.add_run(part)
            set_run(r, size=base_size, bold=base_bold)


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# --------------------------------------------------------------------------- #
#  Styles
# --------------------------------------------------------------------------- #
def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = BODY
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:ascii"), BODY)
    normal.element.rPr.rFonts.set(qn("w:hAnsi"), BODY)
    normal.element.rPr.rFonts.set(qn("w:cs"), BODY)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h_sizes = {"Heading 1": 16, "Heading 2": 13, "Heading 3": 12}
    for name, size in h_sizes.items():
        st = doc.styles[name]
        st.font.name = BODY
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
        st.element.rPr.rFonts.set(qn("w:ascii"), BODY)
        st.element.rPr.rFonts.set(qn("w:hAnsi"), BODY)
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True

    try:
        cap = doc.styles["Caption"]
    except KeyError:
        cap = doc.styles.add_style("Caption", 1)
    cap.font.name = BODY
    cap.font.size = Pt(11)
    cap.font.italic = True
    cap.font.bold = False
    cap.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(12)


# --------------------------------------------------------------------------- #
#  Title / front matter
# --------------------------------------------------------------------------- #
def centered(doc, text, size, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, italic=italic)
    return p


def build_title_page(doc):
    for _ in range(2):
        doc.add_paragraph()
    centered(doc, "Semantic Services Rwanda", 18, bold=True, space_after=4)
    centered(doc, "AI & Recruitment Technology Division", 14, space_after=2)
    centered(doc, "Product & Technical Proposal", 14, space_after=24)

    centered(doc, "AI-Powered Candidate Fraud & Online", 20, bold=True, space_after=2)
    centered(doc, "Assessment Integrity System (SemanticGuard AI)", 20, bold=True, space_after=20)

    centered(
        doc,
        "A Product and Technical Proposal prepared for Semantic Services "
        "Rwanda, outlining the design and delivery of an AI-powered integrity "
        "platform for online recruitment and candidate assessment",
        12,
        italic=True,
        space_after=24,
    )

    centered(doc, "Prepared By", 12, bold=True, space_after=6)

    details = [
        ("Name", "Shingiro Faisal"),
        ("Role", "Software Engineer"),
        ("Phone No", "0787947046"),
        ("Email", "faisalshingiro10@gmail.com"),
        ("Date", "January 2026"),
    ]
    tbl = doc.add_table(rows=len(details), cols=2)
    tbl.alignment = 1  # center
    tbl.autofit = True
    for i, (k, v) in enumerate(details):
        c0, c1 = tbl.rows[i].cells
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rk = c0.paragraphs[0].add_run(f"{k}:")
        set_run(rk, bold=True)
        rv = c1.paragraphs[0].add_run(v)
        set_run(rv)
    page_break(doc)


def build_toc_pages(doc):
    h = doc.add_heading("TABLE OF CONTENTS", level=1)
    h.paragraph_format.page_break_before = False
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u',
              "Right-click and choose 'Update Field' to build the contents.")
    page_break(doc)

    doc.add_heading("LIST OF FIGURES", level=1).paragraph_format.page_break_before = False
    p = doc.add_paragraph()
    add_field(p, 'TOC \\h \\z \\c "Figure"',
              "Right-click and choose 'Update Field' to build the list of figures.")
    doc.add_paragraph()
    doc.add_heading("LIST OF TABLES", level=1).paragraph_format.page_break_before = False
    p = doc.add_paragraph()
    add_field(p, 'TOC \\h \\z \\c "Table"',
              "Right-click and choose 'Update Field' to build the list of tables.")
    page_break(doc)


# --------------------------------------------------------------------------- #
#  Body builders
# --------------------------------------------------------------------------- #
SEQ_COUNTERS = {"Figure": 0, "Table": 0}


def add_image(doc, image_path: Path):
    """Insert an image on its own page, scaled to fit the printable area."""
    page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    pic = run.add_picture(str(image_path), width=MAX_IMG_W)
    # Maintain aspect ratio but never exceed the usable page height.
    if pic.height > MAX_IMG_H:
        ratio = MAX_IMG_H / pic.height
        pic.height = int(pic.height * ratio)
        pic.width = int(pic.width * ratio)


def add_caption(doc, label, title):
    SEQ_COUNTERS[label] = SEQ_COUNTERS.get(label, 0) + 1
    number = SEQ_COUNTERS[label]

    # Figures are illustrated by a real image placed on its own page.
    if label == "Figure":
        image_file = FIGURE_IMAGES.get(title)
        if image_file:
            img_path = IMAGES / image_file
            if img_path.exists():
                add_image(doc, img_path)
            else:
                print(f"WARNING: missing image for '{title}': {img_path}")

    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{label} ")
    set_run(r, size=11, italic=True)
    # SEQ field cached with the correct sequential number so it displays
    # correctly even before the reader presses F9 to update fields.
    add_field(p, f'SEQ {label} \\* ARABIC', str(number))
    r2 = p.add_run(f": {title}")
    set_run(r2, size=11, italic=True)


def add_table(doc, rows):
    header = [c.strip() for c in rows[0]]
    data = rows[1:]
    tbl = doc.add_table(rows=1, cols=len(header))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = 1
    for j, txt in enumerate(header):
        cell = tbl.rows[0].cells[j]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(cell.paragraphs[0], txt)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in data:
        cells = tbl.add_row().cells
        for j, txt in enumerate(row):
            if j < len(cells):
                add_inline(cells[j].paragraphs[0], txt.strip())
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def flush_list(doc, items):
    for item in items:
        indent, kind, text = item[0], item[1], item[2]
        number = item[3] if len(item) > 3 else None
        if kind == "num":
            # Render the literal markdown number so each section restarts at 1
            # rather than continuing the document-wide list counter.
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.left_indent = Inches(0.5)
            pf.first_line_indent = Inches(-0.3)
            rn = p.add_run(f"{number}. ")
            set_run(rn)
            add_inline(p, text)
        else:
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline(p, text)


def parse_table_block(lines, i):
    rows = []
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        raw = lines[i].strip().strip("|")
        cells = [c.strip() for c in raw.split("|")]
        if not all(set(c) <= set("-: ") for c in cells):  # skip separator row
            rows.append(cells)
        i += 1
    return rows, i


def build_body(doc, text):
    lines = text.split("\n")
    i = 0
    current_h1 = ""
    pending_list: list = []

    def flush():
        nonlocal pending_list
        if pending_list:
            flush_list(doc, pending_list)
            pending_list = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # skip separators / page-break divs / blank handling
        if stripped == "---" or stripped.startswith("<div"):
            i += 1
            continue
        if not stripped:
            flush()
            i += 1
            continue

        # diagrams (code fences) — skipped; replaced by real images at caption
        if stripped.startswith("```"):
            flush()
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1  # closing fence
            continue

        # captions
        m = CAPTION_RE.match(stripped)
        if m:
            flush()
            add_caption(doc, m.group(1), m.group(2))
            i += 1
            continue

        # headings
        m = HEADING_RE.match(stripped)
        if m:
            flush()
            level = len(m.group(1))
            title = m.group(2).strip()
            if level == 1:
                current_h1 = title.upper()
                h = doc.add_heading(title, level=1)
                h.paragraph_format.page_break_before = True
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(title, level=min(level, 3))
            i += 1
            continue

        # tables
        if stripped.startswith("|"):
            flush()
            rows, i = parse_table_block(lines, i)
            if rows:
                add_table(doc, rows)
            continue

        # blockquote
        if stripped.startswith(">"):
            flush()
            qtext = stripped.lstrip(">").strip().strip("_")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            add_inline(p, qtext, base_size=11)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # numbered list
        m = NUMBER_RE.match(line)
        if m:
            indent = len(m.group(1))
            pending_list.append((indent, "num", m.group(3).strip(), m.group(2)))
            i += 1
            # gather indented continuation lines
            while i < len(lines) and lines[i].strip() and \
                    not HEADING_RE.match(lines[i].strip()) and \
                    not BULLET_RE.match(lines[i]) and not NUMBER_RE.match(lines[i]) and \
                    lines[i].startswith(" "):
                idx = len(pending_list) - 1
                prev = pending_list[idx]
                pending_list[idx] = (prev[0], "num",
                                     prev[2] + " " + lines[i].strip(), prev[3])
                i += 1
            continue

        # bullet list
        m = BULLET_RE.match(line)
        if m:
            indent = len(m.group(1))
            pending_list.append((indent, "bul", m.group(2).strip()))
            i += 1
            while i < len(lines) and lines[i].strip() and \
                    not HEADING_RE.match(lines[i].strip()) and \
                    not BULLET_RE.match(lines[i]) and not NUMBER_RE.match(lines[i]) and \
                    lines[i].startswith(" "):
                idx = len(pending_list) - 1
                pending_list[idx] = (indent, "bul",
                                     pending_list[idx][2] + " " + lines[i].strip())
                i += 1
            continue

        # plain paragraph
        flush()
        p = doc.add_paragraph()
        if current_h1 == "REFERENCES":
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline(p, stripped)
        i += 1

    flush()


# --------------------------------------------------------------------------- #
#  Main
# --------------------------------------------------------------------------- #
def main():
    md = SRC.read_text(encoding="utf-8")
    marker = "# ABSTRACT"
    idx = md.find(marker)
    body_md = md[idx:]

    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.27)   # A4
        section.page_height = Inches(11.69)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    configure_styles(doc)
    build_title_page(doc)
    build_toc_pages(doc)
    build_body(doc, body_md)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
